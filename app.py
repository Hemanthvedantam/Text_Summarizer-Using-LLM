from flask import Flask, render_template, request, send_file, redirect, url_for, flash, session, jsonify
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
import textstat
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import base64
from PyPDF2 import PdfReader
from docx import Document
import requests
from bs4 import BeautifulSoup
from readability import Document as ReadabilityDocument
import re
import io
import speech_recognition as sr
import os
from fpdf import FPDF
import uuid
from datetime import datetime
import random
import tempfile
from pathlib import Path
import torch
from io import BytesIO
import nltk
from flask_cors import CORS


DATA_DIR = Path(tempfile.gettempdir()) / "quantumtext"
DATA_DIR.mkdir(exist_ok=True)
app = Flask(__name__)
CORS(app)
app.secret_key = os.urandom(24)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10MB limit

nltk.download(['punkt', 'averaged_perceptron_tagger'])
recognizer = sr.Recognizer()

model_name = "google/flan-t5-small"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
summarizer = pipeline(
    task="summarization",
    model=model,
    tokenizer=tokenizer,
    device=-1,  # Force CPU
    # Anti-repetition parameters:
    do_sample=True,          # Add randomness
    temperature=0.9,         # Lower = less random (0.7-1.0 works well)
    repetition_penalty=2.5,  # Penalize repeated phrases (higher = stronger penalty)
    top_p=0.95,              # Nucleus sampling: focus on high-probability tokens
    max_length=100,          # Adjust based on your input text length
    min_length=30,           # Avoid overly short summaries
)

def save_analysis_data(data):
    """Save analysis data to disk and return UUID"""
    analysis_id = str(uuid.uuid4())
    file_path = DATA_DIR / f"{analysis_id}.json"
    
    import json
    with open(file_path, 'w') as f:
        json.dump(data, f)
    
    return analysis_id

def load_analysis_data(analysis_id):
    """Load analysis data from disk"""
    file_path = DATA_DIR / f"{analysis_id}.json"
    
    if not file_path.exists():
        return None
    
    import json
    with open(file_path, 'r') as f:
        return json.load(f)

def process_web_content(url):
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'}
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()

        doc = ReadabilityDocument(response.text)
        content = doc.summary()
        soup = BeautifulSoup(content, 'html.parser')
        text = soup.get_text(separator='\n', strip=True)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip() or None
    except Exception as e:
        flash(f"Error processing website: {str(e)}")
        return None

def process_file(file):
    try:
        if file.content_type == "application/pdf":
            pdf = PdfReader(file)
            text = "\n".join([page.extract_text() or '' for page in pdf.pages])
            return text.strip() or None

        elif file.content_type in ["audio/mpeg", "audio/wav"]:
            with sr.AudioFile(file.stream) as source:
                audio = recognizer.record(source)
                return recognizer.recognize_google(audio)

        elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(BytesIO(file.read()))
            return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        elif file.content_type == "text/plain":
            return file.read().decode(errors='replace')

        else:
            flash("Unsupported file format")
            return None

    except Exception as e:
        flash(f"File processing error: {str(e)}")
        return None

def generate_summary(text, target_words=300, summary_format='paragraph'):
    try:
        max_model_length = 512
        chunks = []
        current_chunk = []
        current_length = 0

        # Add format-specific instruction
        if summary_format == 'bullets':
            instruction = "Summarize the following text in concise bullet points:\n"
        else:
            instruction = "Provide a detailed paragraph summary of the following text:\n"

        sentences = nltk.sent_tokenize(text)
        
        # Account for instruction token length
        instruction_tokens = tokenizer(instruction, return_tensors="pt", truncation=False).input_ids.shape[1]
        max_chunk_tokens = max_model_length - instruction_tokens - 30  # Reserve space for instruction

        for sentence in sentences:
            sentence_tokens = tokenizer(sentence, return_tensors="pt", truncation=False).input_ids.shape[1]
            if current_length + sentence_tokens > max_chunk_tokens:
                chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_length = sentence_tokens
            else:
                current_chunk.append(sentence)
                current_length += sentence_tokens

        if current_chunk:
            chunks.append(' '.join(current_chunk))

        summaries = []
        for chunk in chunks:
            # Prepend instruction to each chunk
            formatted_chunk = instruction + chunk
            inputs = tokenizer(
                formatted_chunk,
                max_length=max_model_length,
                truncation=True,
                return_tensors="pt"
            )
            
            # Generation parameters adjusted for better formatting
            summary_ids = model.generate(
                inputs.input_ids,
                attention_mask=inputs.attention_mask,
                max_length=min(target_words + 50, max_model_length),
                min_length=max(target_words - 50, 10),
                length_penalty=2.0,
                num_beams=4,
                early_stopping=True,
                no_repeat_ngram_size=3  # Reduce repetition
            )
            
            summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
            summaries.append(summary)

        combined_summary = ' '.join(summaries)
        
        # Final refinement pass
        if len(summaries) > 1:
            final_input = instruction + combined_summary
            final_inputs = tokenizer(
                final_input,
                max_length=max_model_length,
                truncation=True,
                return_tensors="pt"
            )
            final_summary_ids = model.generate(
                final_inputs.input_ids,
                attention_mask=final_inputs.attention_mask,
                max_length=target_words + 20,
                min_length=target_words - 20,
                length_penalty=2.0,
                num_beams=4,
                early_stopping=True
            )
            combined_summary = tokenizer.decode(final_summary_ids[0], skip_special_tokens=True)

        # Post-formatting for bullets
        if summary_format == 'bullets':
            # Convert newlines to bullet points
            combined_summary = '\n- '.join(filter(None, combined_summary.split('\n')))
            if not combined_summary.startswith('-'):
                combined_summary = '- ' + combined_summary

        return combined_summary
    except Exception as e:
        flash(f"Summarization error: {str(e)}")
        return None

# Add this function definition BEFORE your route handlers
def get_readability_scores(text):
    """Calculate various readability scores for the given text"""
    try:
        return {
            'flesch': textstat.flesch_reading_ease(text),
            'smog': textstat.smog_index(text),
            'grade': textstat.text_standard(text, float_output=True),
            'coleman_liau': textstat.coleman_liau_index(text),
            'automated_readability': textstat.automated_readability_index(text)
        }
    except Exception as e:
        print(f"Readability calculation error: {str(e)}")
        return {
            'flesch': 0,
            'smog': 0,
            'grade': 'N/A',
            'coleman_liau': 0,
            'automated_readability': 0
        }

def generate_wordcloud(text):
    """Generate a word cloud image from text and return as base64 string"""
    try:
        # Configure word cloud
        wc = WordCloud(
            width=1200,
            height=600,
            background_color='rgba(255, 255, 255, 0)',
            mode='RGBA',
            colormap='plasma',
            max_words=200,
            contour_width=1,
            contour_color='#a855f7',
            stopwords=STOPWORDS
        ).generate(text)

        # Create figure
        plt.figure(figsize=(12, 6))
        plt.imshow(wc, interpolation='bilinear')
        plt.axis('off')
        
        # Save to buffer
        buffer = io.BytesIO()
        plt.savefig(buffer, format='png', transparent=True, bbox_inches='tight')
        plt.close()
        
        # Convert to base64
        buffer.seek(0)
        img_b64 = base64.b64encode(buffer.getvalue()).decode()
        return f"data:image/png;base64,{img_b64}"
    except Exception as e:
        print(f"Word cloud generation error: {str(e)}")
        return None

class RuleBasedAssistant:
    def __init__(self):
        self.rules = [
            (r'(hi|hello|hey|greetings|good day|howdy)\b', self.handle_greeting),
            (r'(bye|goodbye|see ya|exit|quit)\b', self.handle_goodbye),
            (r'(help|support|assistance)\b', self.handle_help),
            (r'(summarize|summary|analyze)\b', self.handle_summary_request),
            (r'(thank you|thanks|appreciate)\b', self.handle_thanks),
            (r'how are you|how\'s it going', self.handle_status),
            (r'what can you do|your purpose', self.handle_capabilities)
        ]
        
        self.context = {}
        self.default_responses = [
            "I'm here to help with text analysis!",
            "Let me know how I can assist with your content.",
            "Would you like me to analyze some text for you?"
        ]

    def match_pattern(self, pattern, text):
        return re.search(pattern, text, re.IGNORECASE)

    def respond(self, user_input):
        user_input = user_input.lower().strip()
        response = None
        
        # Check patterns
        for pattern, handler in self.rules:
            if self.match_pattern(pattern, user_input):
                response = handler(user_input)
                if response: break
        
        # Fallback response
        if not response:
            response = random.choice(self.default_responses)
            
        return {
            'user': user_input,
            'bot': response,
            'timestamp': datetime.now().isoformat()
        }

    # Handler methods
    def handle_greeting(self, text):
        greetings = [
            "Hello! Ready to analyze some text?",
            "Hi there! How can I assist you today?",
            "Greetings! Let's work on some content analysis."
        ]
        return random.choice(greetings)
    
    def handle_goodbye(self, text):
        farewells = [
            "Goodbye! Come back with more text to analyze!",
            "See you later! Don't hesitate to return with documents.",
            "Farewell! Remember I'm here for all your analysis needs."
        ]
        return random.choice(farewells)
    
    def handle_help(self, text):
        return ("I can help you with:\n"
                "- Text summarization\n"
                "- Readability analysis\n"
                "- File processing (PDF, DOCX, TXT)\n"
                "- Web content analysis\n"
                "Just upload or paste your content to get started!")
    
    def handle_summary_request(self, text):
        return "Sure! Please paste your text, upload a file, or provide a URL to get started."
    
    def handle_thanks(self, text):
        return "You're welcome! Let me know if you need anything else."
    
    def handle_status(self, text):
        return "I'm functioning optimally and ready to analyze text!"
    
    def handle_capabilities(self, text):
        return ("I specialize in:\n"
                "- AI-powered text summarization\n"
                "- Readability scoring\n"
                "- File format conversion\n"
                "- Content analysis visualization\n"
                "Try pasting some text to see my capabilities!")

# Initialize assistant
assistant = RuleBasedAssistant()

@app.route('/chat', methods=['POST'])
def handle_chat():
    try:
        user_message = request.json.get('message', '')
        if not user_message.strip():
            return jsonify({'error': 'Empty message'}), 400
        
        response = assistant.respond(user_message)
        # Ensure proper JSON formatting
        return jsonify({
            'user': response['user'],
            'bot': response['bot'],
            'timestamp': response['timestamp']
        })
    
    except Exception as e:
        app.logger.error(f"Chat error: {str(e)}")
        return jsonify({
            'error': f'Chat error: {str(e)}'
        }), 500


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        session.clear()
        input_method = request.form.get('input_method', 'text')
        text = ""
        summary_length = int(request.form.get('summary_length', 300))
        summary_format = request.form.get('summary_format', 'paragraph')
        summary = generate_summary(text, summary_length, summary_format)
        
        try:
            if input_method == "text":
                text = request.form.get('input_text', '').strip()
                if not text:
                    flash("Please enter some text to analyze")
                    return redirect(url_for('index'))

            elif input_method == "file":
                if 'file' not in request.files:
                    flash("No file selected")
                    return redirect(url_for('index'))
                
                file = request.files['file']
                if file.filename == '':
                    flash("No file selected")
                    return redirect(url_for('index'))
                
                text = process_file(file)
                if not text:
                    return redirect(url_for('index'))

            elif input_method == "url":
                url = request.form.get('url', '').strip()
                if not url:
                    flash("Please enter a URL")
                    return redirect(url_for('index'))
                
                text = process_web_content(url)
                if not text:
                    return redirect(url_for('index'))

            word_count = textstat.lexicon_count(text, removepunct=True)
            if word_count < 100:
                flash(f"Minimum 100 words required (found {word_count})")
                return redirect(url_for('index'))

            summary = generate_summary(text, summary_length,summary_format)
            if not summary:
                return redirect(url_for('index'))

            analysis_data = {
                'summary': summary,
                'readability': get_readability_scores(text),
                'word_count': word_count
            }
            
            analysis_id = save_analysis_data(analysis_data)
            session['analysis_id'] = analysis_id
            
            return redirect(url_for('results'))

        except Exception as e:
            flash(f"Processing error: {str(e)}")
            return redirect(url_for('index'))

    return render_template('index.html')

@app.route('/results')
def results():
    analysis_id = session.get('analysis_id')
    if not analysis_id:
        return redirect(url_for('index'))
    
    data = load_analysis_data(analysis_id)
    if not data:
        flash("Analysis expired or not found")
        return redirect(url_for('index'))
    
    return render_template('index.html',
                         summary=data['summary'],
                         readability=data['readability'],
                         word_count=data['word_count'])

@app.route('/download', methods=['POST'])
def download():
    analysis_id = session.get('analysis_id')
    if not analysis_id:
        flash("No analysis found")
        return redirect(url_for('index'))
    
    data = load_analysis_data(analysis_id)
    if not data:
        flash("Analysis expired or not found")
        return redirect(url_for('index'))
    
    summary = data.get('summary', '')
    format_type = request.form.get('format')
    
    if format_type == 'text':
        return send_file(
            io.BytesIO(summary.encode()),
            mimetype='text/plain',
            as_attachment=True,
            download_name='summary.txt'
        )
    elif format_type == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, summary)
        return send_file(
            io.BytesIO(pdf.output()),
            mimetype='application/pdf',
            as_attachment=True,
            download_name='summary.pdf'
        )
    elif format_type == 'word':
        doc = Document()
        for line in summary.split('\n'):
            doc.add_paragraph(line)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='summary.docx'
        )
    
    return redirect(url_for('results'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
